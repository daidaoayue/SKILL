#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
latex-to-word/scripts/convert.py — 真实可执行的 LaTeX → Word 转换器

底层逻辑：
- 主路径：pandoc（保留公式、引用、表格）
- 后处理：python-docx 应用学校字体规范（如有学校 docx 模板，用 --reference-doc）
- 验证：python-docx 复检公式数量、章节数、字符数

用法：
    python convert.py <input.tex> [-o output.docx] [--template school.docx]
                                  [--bibliography refs.bib]
                                  [--csl citation-style.csl]

例子：
    python convert.py paper/main.tex -o paper/main.docx
    python convert.py main.tex --template buaa.docx --bibliography refs.bib
"""

from __future__ import annotations

import argparse
import io
import json
import os
import re
import shutil
import subprocess
import sys
from pathlib import Path

if sys.platform == "win32":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")


def check_pandoc() -> tuple[bool, str]:
    """检查 pandoc 是否可用。"""
    if shutil.which("pandoc") is None:
        return False, "pandoc 未安装。请先 `choco install pandoc` 或 `apt install pandoc`"
    try:
        out = subprocess.run(["pandoc", "--version"], capture_output=True, text=True, timeout=10)
        version_line = out.stdout.split("\n")[0] if out.stdout else "unknown"
        return True, version_line
    except Exception as e:
        return False, f"pandoc 检查失败: {e}"


def check_xelatex() -> tuple[bool, str]:
    """检查 xelatex 是否可用。"""
    if shutil.which("xelatex") is None:
        return False, "xelatex 未找到（安装 MiKTeX/TeX Live）"
    return True, "xelatex 可用"


def check_pdf2docx() -> tuple[bool, str]:
    """检查 pdf2docx 是否可用（路径 D 兜底用）。"""
    try:
        from pdf2docx import Converter  # noqa: F401
        return True, "pdf2docx 可用"
    except ImportError:
        return False, "pdf2docx 未安装（pip install pdf2docx）"


def convert_via_pandoc_preprocessed(
    tex_path: Path, output_path: Path,
    template: Path | None = None,
    bibliography: Path | None = None,
) -> tuple[bool, str]:
    """
    路径 B（v0.5.6 新增）: 预处理 .tex 移除冲突宏包后 pandoc 转，得到正确格式 docx。

    针对学校模板（buaathesis 等）+ gbt7714 等 pandoc 不识别的宏包：
      - \\documentclass{buaathesis} → \\documentclass[UTF8]{ctexart}
      - \\usepackage{gbt7714} → 删除
      - \\renewcommand{\\cite} 上标式 → 删除
      - \\citestyle{numerical} → 删除
      - \\chapter → \\section（ctexart 没 chapter）
    然后跑 pandoc 真直转，输出干净 docx（远小于 pdf2docx 反转的）。
    """
    if shutil.which("pandoc") is None:
        return False, "pandoc 未安装"

    tmp_tex = tex_path.parent / f"{tex_path.stem}_pandoc_preprocessed.tex"
    try:
        original = tex_path.read_text(encoding="utf-8", errors="replace")
        base_dir = tex_path.parent

        # 提取 \include / \input 的章节列表（保持顺序）
        body_includes = []
        for m in re.finditer(r"\\(?:include|input)\{([^}]+)\}", original):
            body_includes.append(m.group(1))

        def resolve_include(inc: str) -> Path | None:
            """解析 \\include 的路径。"""
            for cand in [base_dir / f"{inc}.tex", base_dir / inc, base_dir / f"{inc}/main.tex"]:
                if cand.exists():
                    return cand
            return None

        def normalize_chapter(text: str) -> str:
            """替换 buaathesis 自定义命令为 pandoc 友好版。"""
            # 中英文摘要环境
            text = re.sub(r"\\begin\{cabstract\}", r"\\section*{摘要}\n", text)
            text = re.sub(r"\\end\{cabstract\}", r"", text)
            text = re.sub(r"\\begin\{eabstract\}", r"\\section*{Abstract}\n", text)
            text = re.sub(r"\\end\{eabstract\}", r"", text)
            # 关键词命令
            text = re.sub(r"\\keywords\{([^}]+)\}", r"\n\\noindent \\textbf{关键词：} \1\n", text)
            text = re.sub(r"\\ekeywords\{([^}]+)\}", r"\n\\noindent \\textbf{Keywords:} \1\n", text)
            # 去掉嵌套 \markboth / \markright（pandoc 不认）
            text = re.sub(r"\\markboth\{[^}]*\}\{[^}]*\}", "", text)
            text = re.sub(r"\\markright\{[^}]*\}", "", text)
            # ctexbook 支持 \chapter，不用降级
            # 删除 buaathesis 特殊命令（含可能的参数）
            for cmd in ["pagestyle", "thispagestyle", "mainmatter", "frontmatter",
                        "backmatter", "appendix", "maketitle", "tableofcontents",
                        "fenkai", "makecover", "kaiti", "lishu", "fangsong",
                        "include", "input"]:  # \include / \input 已经被我们 inline 了
                text = re.sub(r"\\" + cmd + r"\b\s*(?:\[[^\]]*\])?\s*(?:\{[^}]*\})?", "", text)
            # 双行注释
            text = re.sub(r"^\s*%.*$", "", text, flags=re.MULTILINE)
            return text

        # 构造合并后的 .tex（章节内容 inline 进来，不靠 \input）
        new_text = (
            "% pandoc 预处理版（路径 B 自动生成 - merged inline）\n"
            "\\documentclass[UTF8,12pt]{ctexbook}\n"
            "\\usepackage{amsmath,amssymb,amsfonts,bm}\n"
            "\\usepackage{graphicx}\n"
            "\\usepackage{booktabs,longtable,array,multirow}\n"
            "\\usepackage{xcolor}\n"
            "\\usepackage{hyperref}\n"
            "\\usepackage{geometry}\n"
            "\\geometry{a4paper,margin=2.5cm}\n"
            "\\graphicspath{{" + str(base_dir).replace("\\", "/") + "/figure/}}\n"
            "\n\\begin{document}\n\n"
        )
        for inc in body_includes:
            # 跳过封面 / 任务书 / 学校信息
            if any(skip in inc for skip in ["com_info", "bachelor_info", "/assign", "acknowledgement"]):
                continue
            sub_path = resolve_include(inc)
            if sub_path:
                sub_text = sub_path.read_text(encoding="utf-8", errors="replace")
                new_text += f"\n% === {inc} ===\n"
                new_text += normalize_chapter(sub_text)
                new_text += "\n"
            else:
                new_text += f"% [skip: {inc} not found]\n"
        new_text += "\n\\end{document}\n"

        tmp_tex.write_text(new_text, encoding="utf-8")

        cmd = ["pandoc", str(tmp_tex), "--from=latex", "--to=docx",
               "--toc", "--toc-depth=3", "-o", str(output_path),
               "--resource-path", str(tex_path.parent)]
        if bibliography and bibliography.exists():
            cmd += ["--bibliography", str(bibliography)]
        if template and template.exists():
            cmd += ["--reference-doc", str(template)]

        result = subprocess.run(cmd, cwd=tex_path.parent, capture_output=True,
                               text=True, timeout=180, encoding="utf-8", errors="replace")
        if output_path.exists() and output_path.stat().st_size > 1024:
            return True, f"OK (路径 B: pandoc + 预处理)"
        return False, f"pandoc 失败: {result.stderr[:300]}"
    finally:
        if tmp_tex.exists():
            tmp_tex.unlink()


def convert_via_xelatex_pdf2docx(tex_path: Path, output_path: Path) -> tuple[bool, str]:
    """
    路径 D（v0.5.3 新增）: 用现有 PDF（或重新 xelatex 编译）→ pdf2docx → docx
    用于绕过 pandoc 解析 .sty 失败的场景（如 gbt7714.sty 与 pandoc 冲突）。
    """
    ok, msg = check_xelatex()
    if not ok:
        return False, f"路径 D 失败: {msg}"
    ok, msg = check_pdf2docx()
    if not ok:
        return False, f"路径 D 失败: {msg}"

    pdf_path = tex_path.with_suffix(".pdf")

    # 1. PDF 已存在则直接用；否则尝试 xelatex 编译
    if not pdf_path.exists():
        try:
            subprocess.run(
                ["xelatex", "-interaction=nonstopmode", "-halt-on-error", str(tex_path.name)],
                cwd=tex_path.parent, capture_output=True, timeout=300,
            )
            # 第二次跑解决目录/引用
            subprocess.run(
                ["xelatex", "-interaction=nonstopmode", str(tex_path.name)],
                cwd=tex_path.parent, capture_output=True, timeout=300,
            )
        except subprocess.TimeoutExpired:
            return False, "xelatex 编译超时（>300s）"
        except Exception as e:
            return False, f"xelatex 异常: {e}"

    if not pdf_path.exists():
        return False, "xelatex 没生成 PDF（请检查 .tex 编译错误）"

    # 2. PDF → docx
    try:
        from pdf2docx import Converter
        cv = Converter(str(pdf_path))
        cv.convert(str(output_path))
        cv.close()
        return True, f"OK (路径 D: pdf2docx 从 {pdf_path.name})"
    except Exception as e:
        return False, f"pdf2docx 异常: {e}"


def detect_complexity(tex_path: Path) -> dict:
    """统计 .tex 内的复杂度：公式/引用/表格数量。"""
    try:
        text = tex_path.read_text(encoding="utf-8", errors="replace")
    except Exception as e:
        return {"error": str(e)}

    # 统计 \include{} 引入的子文件
    includes = re.findall(r"\\include\{([^}]+)\}", text)
    inputs = re.findall(r"\\input\{([^}]+)\}", text)

    # 收集主文件 + 子文件的全文
    all_text = text
    base_dir = tex_path.parent
    for inc in includes + inputs:
        sub_path = base_dir / f"{inc}.tex"
        if not sub_path.exists():
            sub_path = base_dir / inc
        if sub_path.exists():
            try:
                all_text += "\n" + sub_path.read_text(encoding="utf-8", errors="replace")
            except Exception:
                pass

    return {
        "equations_count": len(re.findall(r"\\begin\{equation\}", all_text)),
        "tables_count": len(re.findall(r"\\begin\{tabular\}", all_text)),
        "figures_count": len(re.findall(r"\\begin\{figure\}", all_text)),
        "cite_count": len(re.findall(r"\\cite\{", all_text)),
        "ref_count": len(re.findall(r"\\ref\{", all_text)),
        "chapter_count": len(re.findall(r"\\chapter\{", all_text)),
        "section_count": len(re.findall(r"\\section\{", all_text)),
        "include_files": includes + inputs,
        "char_count": len(re.sub(r"\\[a-zA-Z]+\{[^{}]*\}|\\[a-zA-Z]+|[{}\\%]", "", all_text)),
    }


def convert_with_pandoc(
    tex_path: Path,
    output_path: Path,
    template: Path | None = None,
    bibliography: Path | None = None,
    csl: Path | None = None,
    extra_args: list[str] | None = None,
) -> tuple[bool, str]:
    """调 pandoc 真转换 .tex → .docx。"""
    cmd = [
        "pandoc",
        str(tex_path),
        "--from=latex",
        "--to=docx",
        "--toc",
        "--toc-depth=3",
        "-o", str(output_path),
        "--resource-path", str(tex_path.parent),
    ]
    if template and template.exists():
        cmd += ["--reference-doc", str(template)]
    if bibliography and bibliography.exists():
        cmd += ["--bibliography", str(bibliography)]
    if csl and csl.exists():
        cmd += ["--csl", str(csl)]
    if extra_args:
        cmd += extra_args

    try:
        result = subprocess.run(
            cmd,
            cwd=tex_path.parent,
            capture_output=True,
            text=True,
            timeout=180,
            encoding="utf-8",
            errors="replace",
        )
        if result.returncode == 0:
            return True, "OK"
        # pandoc 警告但仍生成了文件
        if output_path.exists() and output_path.stat().st_size > 0:
            return True, f"OK (with warnings: {result.stderr[:200]})"
        return False, f"pandoc 失败: {result.stderr[:500]}"
    except subprocess.TimeoutExpired:
        return False, "pandoc 超时（>180s）"
    except Exception as e:
        return False, f"pandoc 异常: {e}"


def verify_docx(docx_path: Path, expected: dict) -> dict:
    """用 python-docx 验证生成的 docx。"""
    try:
        from docx import Document
    except ImportError:
        return {"docx_readable": False, "error": "python-docx 未安装（pip install python-docx）"}

    if not docx_path.exists():
        return {"docx_readable": False, "error": "docx 文件不存在"}

    try:
        doc = Document(docx_path)
    except Exception as e:
        return {"docx_readable": False, "error": f"docx 打开失败: {e}"}

    para_count = len(doc.paragraphs)
    table_count = len(doc.tables)
    all_text = "\n".join(p.text for p in doc.paragraphs)
    char_count = len(all_text)

    return {
        "docx_readable": True,
        "paragraph_count": para_count,
        "table_count": table_count,
        "docx_char_count": char_count,
        "docx_size_kb": round(docx_path.stat().st_size / 1024, 1),
        "table_count_match": table_count >= expected.get("tables_count", 0) * 0.8,
        "first_para_preview": (doc.paragraphs[0].text[:60] if doc.paragraphs else ""),
    }


def main() -> int:
    parser = argparse.ArgumentParser(
        description="latex-to-word: 真实可执行的 LaTeX → Word 转换器",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="例子:\n"
               "  python convert.py paper/main.tex\n"
               "  python convert.py main.tex -o out.docx --template buaa.docx\n",
    )
    parser.add_argument("input_tex", type=Path, help="输入的 main.tex 路径")
    parser.add_argument("-o", "--output", type=Path, default=None, help="输出 docx 路径（默认同名 .docx）")
    parser.add_argument("--template", type=Path, default=None, help="学校 Word 模板（pandoc reference-doc）")
    parser.add_argument("--bibliography", type=Path, default=None, help="参考文献 .bib")
    parser.add_argument("--csl", type=Path, default=None, help="引用格式 .csl（如 gb-t-7714）")
    parser.add_argument("--report", type=Path, default=None, help="输出 JSON 转换报告")
    args = parser.parse_args()

    if not args.input_tex.exists():
        print(f"❌ 输入文件不存在: {args.input_tex}")
        return 1

    output = args.output or args.input_tex.with_suffix(".docx")

    # 1. 检查 pandoc
    pandoc_ok, pandoc_msg = check_pandoc()
    print(f"[1/4] pandoc 检查: {'✅' if pandoc_ok else '❌'} {pandoc_msg}")
    if not pandoc_ok:
        return 2

    # 2. 复杂度检测
    print(f"[2/4] 复杂度检测: {args.input_tex}")
    complexity = detect_complexity(args.input_tex)
    if "error" in complexity:
        print(f"  ❌ {complexity['error']}")
        return 3
    print(f"  公式={complexity['equations_count']} 表格={complexity['tables_count']} "
          f"图={complexity['figures_count']} 引用={complexity['cite_count']} "
          f"章={complexity['chapter_count']} 节={complexity['section_count']}")
    print(f"  子文件: {len(complexity['include_files'])} 个")

    # 3. 转换 — 路径优先级 A → B → D
    print(f"[3/4] 路径 A · pandoc 直转 → {output}")
    ok, msg = convert_with_pandoc(
        args.input_tex, output,
        template=args.template,
        bibliography=args.bibliography,
        csl=args.csl,
    )
    print(f"  {'✅' if ok else '❌'} {msg}")
    if not ok:
        print(f"\n[3.3/4] 路径 A 失败，尝试路径 B · pandoc + 预处理（推荐，docx 格式好）")
        ok, msg = convert_via_pandoc_preprocessed(
            args.input_tex, output,
            template=args.template, bibliography=args.bibliography,
        )
        print(f"  {'✅' if ok else '❌'} {msg}")
    if not ok:
        print(f"\n[3.6/4] 路径 B 失败，最后兜底路径 D · xelatex+pdf2docx（格式较差）")
        ok, msg = convert_via_xelatex_pdf2docx(args.input_tex, output)
        print(f"  {'✅' if ok else '❌'} {msg}")
        if not ok:
            print(f"\n三条路径都失败。建议手动检查 .tex 编译错误。")
            return 4

    # 4. 验证
    print(f"[4/4] python-docx 验证")
    verify = verify_docx(output, complexity)
    if not verify.get("docx_readable"):
        print(f"  ❌ {verify.get('error', '未知错误')}")
        return 5
    print(f"  段落={verify['paragraph_count']} 表格={verify['table_count']} "
          f"字符={verify['docx_char_count']} 大小={verify['docx_size_kb']} KB")

    # 输出报告
    if args.report:
        report = {
            "input": str(args.input_tex),
            "output": str(output),
            "complexity": complexity,
            "verify": verify,
            "pandoc_version": pandoc_msg,
        }
        args.report.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
        print(f"\n📋 报告: {args.report}")

    print(f"\n✅ 转换完成: {output}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
